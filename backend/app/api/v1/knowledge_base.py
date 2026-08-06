"""
通维AI项目管理系统 - 知识库API
提供知识库CRUD、文档上传、搜索等功能

[PMBOK KA: 沟通管理 | PG: 收尾 (Communications/Closing) — 知识管理、经验库]
对应PMI第6版标准：经验教训、组织过程资产

[CPMAI Phase: CPMAI Phase: Data Understanding | Domain: Data for AI — AI知识库数据理解]"""

from typing import List, Optional
import logging, os, uuid as _uuid
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form, status
from fastapi.responses import FileResponse, Response
import urllib.parse

# 文件扩展名 -> MIME 映射，用于在线预览时确定正确 Content-Type
_EXT_MIME = {
    "pdf": "application/pdf",
    "png": "image/png",
    "jpg": "image/jpeg",
    "jpeg": "image/jpeg",
    "gif": "image/gif",
    "webp": "image/webp",
    "bmp": "image/bmp",
    "svg": "image/svg+xml",
    "txt": "text/plain",
    "md": "text/markdown",
    "csv": "text/csv",
    "json": "application/json",
    "html": "text/html",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "pptx": "application/vnd.openxmlformats-officedocument.presentationml.presentation",
}

from pydantic import BaseModel
from sqlalchemy import select, and_, or_, desc, delete
from sqlalchemy.orm import selectinload
from sqlalchemy.ext.asyncio import AsyncSession

from app.db.session import get_db
from app.models.knowledge_base import KnowledgeBase, KnowledgeDocument, KnowledgeChunk, DocumentStatus, KnowledgeBaseShare, ShareType, SharePermission, UserGroupMember
from app.services.rag_engine import RAGEngine, get_rag_engine
from app.services.kb_access import _can_access_kb, get_accessible_kb_ids
from app.core.ai_engine import ai_engine
from app.core.security import get_current_user
from app.models import User
from app.config import settings

router = APIRouter()

logger = logging.getLogger(__name__)


async def _auth_kb(db: AsyncSession, kb_id: str, user: User, write: bool = False):
    """鉴权：返回 kb 实例；不存在 404，无权 403。"""
    kb = await _can_access_kb(db, kb_id, user, require_write=write)
    if kb is None:
        raise HTTPException(status_code=404, detail="知识库不存在")
    if kb is False:
        raise HTTPException(status_code=403, detail="无权访问该知识库")
    return kb


async def _enrich_kb(kb, db: AsyncSession, user: User) -> dict:
    """返回知识库 dict，并补齐 access_role / visibility / is_shared / my_permission / project_name 字段。"""
    d = kb.to_dict()
    # 关联项目名（如有）：通过 Project 关系查 name
    if getattr(kb, "project", None) is not None and getattr(kb.project, "name", None):
        d["project_name"] = kb.project.name
    elif kb.project_id:
        try:
            from app.models import Project
            proj = await db.get(Project, kb.project_id)
            d["project_name"] = proj.name if proj else None
        except Exception:
            d["project_name"] = None
    else:
        d["project_name"] = None
    if kb.created_by == user.id:
        d["access_role"], d["visibility"], d["is_shared"], d["my_permission"] = "owner", "private", False, "owner"
        return d
    # 系统共享
    sys_row = (
        await db.execute(
            select(KnowledgeBaseShare.id).where(
                KnowledgeBaseShare.kb_id == kb.id,
                KnowledgeBaseShare.share_type == ShareType.SYSTEM.value,
            )
        )
    ).first()
    if sys_row:
        d["access_role"], d["visibility"], d["is_shared"], d["my_permission"] = "system", "system", True, "read"
        return d
    # 用户 / 用户组共享
    shares = (
        await db.execute(select(KnowledgeBaseShare).where(KnowledgeBaseShare.kb_id == kb.id))
    ).scalars().all()
    my_groups = set(
        r[0] for r in (
            await db.execute(select(UserGroupMember.group_id).where(UserGroupMember.user_id == user.id))
        ).all()
    )
    perm = "read"
    for s in shares:
        if s.share_type == ShareType.USER.value and s.target_id == user.id:
            perm = s.permission  # read / write
        elif s.share_type == ShareType.GROUP.value and s.target_id in my_groups:
            perm = s.permission
    d["access_role"], d["visibility"], d["is_shared"], d["my_permission"] = "shared", "shared", True, perm
    return d


async def _answer_with_context(question: str, context: str) -> Optional[str]:
    """QA 的 LLM 回答。

    注意：LLM 调用已暂时禁用，因为服务器上所有 provider 均有连接挂死风险。
    短期内直接返回 None → QA 端点返回原始资料片段作为上下文。
    待配置可靠的 LLM provider（如 MiniMax 已验证可在 20s 内响应）后可恢复调用。
    """
    return None


# ============== Pydantic Schemas ==============

class KnowledgeBaseCreate(BaseModel):
    name: str
    description: Optional[str] = None
    project_id: Optional[str] = None
    embedding_model: Optional[str] = "text-embedding-3-small"
    visibility: Optional[str] = "private"  # private=仅自己 | public=全系统可检索/AI生成可选


class KnowledgeBaseUpdate(BaseModel):
    name: Optional[str] = None
    description: Optional[str] = None
    embedding_model: Optional[str] = None
    # 支持把 KB 关联到具体项目（也可置 null 取消关联）
    project_id: Optional[str] = None
    visibility: Optional[str] = None


class KnowledgeBaseResponse(BaseModel):
    id: str
    name: str
    description: Optional[str]
    project_id: Optional[str]
    embedding_model: str
    created_by: str
    created_at: Optional[str]
    updated_at: Optional[str]
    document_count: int = 0
    access_role: str = "owner"   # owner | shared | system
    is_shared: bool = False
    visibility: str = "private"   # private | shared | system
    my_permission: str = "read"   # owner | write | read（当前用户对该库的权限）

    class Config:
        from_attributes = True


class DocumentCreate(BaseModel):
    title: str
    content: str
    source_type: str = "text"
    source_url: Optional[str] = None
    meta_data: Optional[dict] = None


class DocumentResponse(BaseModel):
    id: str
    kb_id: str
    title: str
    content: Optional[str] = None
    source_type: str
    source_url: Optional[str]
    file_path: Optional[str] = None
    file_name: Optional[str]
    file_size: int
    mime_type: Optional[str]
    status: str
    chunk_count: int
    meta_data: Optional[dict]
    created_at: Optional[str]
    updated_at: Optional[str]

    class Config:
        from_attributes = True


class DocumentUpdate(BaseModel):
    title: Optional[str] = None
    content: Optional[str] = None
    meta_data: Optional[dict] = None


class SearchQuery(BaseModel):
    query: str
    top_k: int = 5


class SearchResult(BaseModel):
    chunk_id: str
    document_id: str
    document_title: str
    content: str
    score: float
    chunk_index: int
    search_method: str


class SearchResponse(BaseModel):
    query: str
    results: List[SearchResult]
    total: int


class RAGQuery(BaseModel):
    query: str
    kb_ids: List[str]
    top_k: int = 5


class RAGResponse(BaseModel):
    query: str
    context: str


# ============== Knowledge Base Endpoints ==============

@router.get("/knowledge-bases/", response_model=List[KnowledgeBaseResponse])
async def list_knowledge_bases(
    project_id: Optional[str] = None,
    scope: str = "all",   # mine=仅自己 | all=自己+分享给我的+系统共享
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    # 可见的 kb id 集合（owner / shared / system）
    accessible = set(await get_accessible_kb_ids(db, current_user, scope="all"))
    owned = set(await get_accessible_kb_ids(db, current_user, scope="mine"))
    system_ids = set(
        r[0] for r in (
            await db.execute(
                select(KnowledgeBaseShare.kb_id).where(KnowledgeBaseShare.share_type == "system")
            )
        ).all()
    )
    shared_ids = accessible - owned - system_ids

    if scope == "mine":
        target_ids = owned
    else:
        target_ids = accessible

    # 一次性取出目标库的分享记录与"我所在的用户组"，用于计算每库的 my_permission
    share_rows = (
        await db.execute(
            select(KnowledgeBaseShare).where(KnowledgeBaseShare.kb_id.in_(target_ids))
        )
    ).scalars().all()
    shares_by_kb: dict = {}
    for s in share_rows:
        shares_by_kb.setdefault(s.kb_id, []).append(s)
    my_groups = set(
        r[0] for r in (
            await db.execute(
                select(UserGroupMember.group_id).where(UserGroupMember.user_id == current_user.id)
            )
        ).all()
    )

    def compute_permission(kb_id: str) -> str:
        if kb_id in owned:
            return "owner"
        if kb_id in system_ids:
            return "read"  # 系统共享恒为只读
        perm = "read"
        for s in shares_by_kb.get(kb_id, []):
            if s.share_type == ShareType.USER.value and s.target_id == current_user.id:
                if s.permission == SharePermission.WRITE.value:
                    return "write"
                perm = "read"
            elif s.share_type == ShareType.GROUP.value and s.target_id in my_groups:
                if s.permission == SharePermission.WRITE.value:
                    return "write"
                perm = "read"
        return perm

    query = select(KnowledgeBase).where(KnowledgeBase.id.in_(target_ids)).options(selectinload(KnowledgeBase.documents))
    if project_id:
        query = query.where(KnowledgeBase.project_id == project_id)
    query = query.order_by(desc(KnowledgeBase.created_at))

    result = await db.execute(query)
    kbs = result.scalars().all()

    out = []
    # 取出本批 KB 涉及的 project_id，批量查 project 名称，避免 N+1
    proj_ids = {k.project_id for k in kbs if k.project_id}
    proj_name_map: dict = {}
    if proj_ids:
        from app.models import Project
        proj_rows = (await db.execute(select(Project).where(Project.id.in_(proj_ids)))).scalars().all()
        proj_name_map = {p.id: p.name for p in proj_rows}

    for kb in kbs:
        d = kb.to_dict()
        d["project_name"] = proj_name_map.get(kb.project_id) if kb.project_id else None
        if kb.id in owned:
            d["access_role"], d["visibility"], d["is_shared"] = "owner", "private", False
        elif kb.id in system_ids:
            d["access_role"], d["visibility"], d["is_shared"] = "system", "system", True
        else:
            d["access_role"], d["visibility"], d["is_shared"] = "shared", "shared", True
        d["my_permission"] = compute_permission(kb.id)
        out.append(d)
    return out


@router.get("/knowledge-bases/ai-selectable", response_model=List[KnowledgeBaseResponse])
async def list_ai_selectable_bases(
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """返回可用于 AI 生成的知识库列表：
       - public（公开）：全系统可见，任何用户可选
       - private（私密）：仅创建者自己可见和可选
       前端用此接口渲染「AI生成时选择知识库」的下拉框（单选）。
    """
    from app.models.knowledge_base import Visibility
    result = await db.execute(
        select(KnowledgeBase).where(
            KnowledgeBase.id.in_(
                # 公开库 + 自己的私密库
                select(KnowledgeBase.id).where(
                    or_(
                        KnowledgeBase.visibility == Visibility.PUBLIC.value,
                        and_(
                            KnowledgeBase.visibility == Visibility.PRIVATE.value,
                            KnowledgeBase.created_by == current_user.id,
                        ),
                    )
                )
            )
        ).order_by(desc(KnowledgeBase.updated_at))
    )
    kbs = result.scalars().all()
    return [kb.to_dict() for kb in kbs]


@router.post("/knowledge-bases/", response_model=KnowledgeBaseResponse, status_code=status.HTTP_201_CREATED)
async def create_knowledge_base(
    data: KnowledgeBaseCreate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    kb = KnowledgeBase(
        name=data.name,
        description=data.description,
        visibility=data.visibility or "private",
        project_id=data.project_id,
        embedding_model=data.embedding_model or "text-embedding-3-small",
        created_by=current_user.id,
    )
    db.add(kb)
    await db.commit()
    await db.refresh(kb)

    return await _enrich_kb(kb, db, current_user)


@router.get("/knowledge-bases/{kb_id}", response_model=KnowledgeBaseResponse)
async def get_knowledge_base(
    kb_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    kb = await _auth_kb(db, kb_id, current_user)
    return await _enrich_kb(kb, db, current_user)


@router.put("/knowledge-bases/{kb_id}", response_model=KnowledgeBaseResponse)
async def update_knowledge_base(
    kb_id: str,
    data: KnowledgeBaseUpdate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    kb = await _auth_kb(db, kb_id, current_user, write=True)

    if data.name is not None:
        kb.name = data.name
    if data.description is not None:
        kb.description = data.description
    if data.embedding_model is not None:
        kb.embedding_model = data.embedding_model
    if "project_id" in data.model_fields_set:
        # 显式置 null 表示解除关联
        kb.project_id = data.project_id
    if data.visibility is not None:
        kb.visibility = data.visibility

    await db.commit()
    await db.refresh(kb)

    return await _enrich_kb(kb, db, current_user)


@router.delete("/knowledge-bases/{kb_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_knowledge_base(
    kb_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    kb = await _auth_kb(db, kb_id, current_user, write=True)

    # 显式级联删除子表，避免生产库（PostgreSQL）外键约束导致 500。
    # 模型 FK 未配置 ondelete=CASCADE，ORM delete-orphan 仅在集合已加载时生效，
    # 因此这里先按 kb_id 批量删除 分享 / 文档块 / 文档，再删除知识库本体。
    await db.execute(
        delete(KnowledgeBaseShare).where(KnowledgeBaseShare.kb_id == kb_id)
    )
    doc_ids = (
        await db.execute(
            select(KnowledgeDocument.id).where(KnowledgeDocument.kb_id == kb_id)
        )
    ).scalars().all()
    if doc_ids:
        await db.execute(
            delete(KnowledgeChunk).where(KnowledgeChunk.document_id.in_(doc_ids))
        )
        await db.execute(
            delete(KnowledgeDocument).where(KnowledgeDocument.kb_id == kb_id)
        )

    await db.delete(kb)
    await db.commit()

    return None


# ============== Document Endpoints ==============

@router.get("/knowledge-bases/{kb_id}/documents", response_model=List[DocumentResponse])
async def list_documents(
    kb_id: str,
    status_filter: Optional[str] = None,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    # 验证知识库权限
    await _auth_kb(db, kb_id, current_user)

    query = select(KnowledgeDocument).where(KnowledgeDocument.kb_id == kb_id)

    if status_filter:
        query = query.where(KnowledgeDocument.status == status_filter)

    query = query.order_by(desc(KnowledgeDocument.created_at))

    result = await db.execute(query)
    docs = result.scalars().all()

    return [doc.to_dict() for doc in docs]


@router.post("/knowledge-bases/{kb_id}/documents", response_model=DocumentResponse, status_code=status.HTTP_201_CREATED)
async def add_document(
    kb_id: str,
    data: DocumentCreate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    # 验证知识库写权限
    await _auth_kb(db, kb_id, current_user, write=True)

    rag_engine = RAGEngine(db_session=db)

    try:
        document = await rag_engine.add_document(
            kb_id=kb_id,
            title=data.title,
            content=data.content,
            source_type=data.source_type,
            source_url=data.source_url,
            meta_data=data.meta_data,
        )
        return document.to_dict()
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"文档处理失败: {str(e)}")


@router.post("/knowledge-bases/{kb_id}/documents/upload", response_model=DocumentResponse)
async def upload_document(
    kb_id: str,
    file: UploadFile = File(...),
    title: Optional[str] = Form(None),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    # 验证知识库写权限
    await _auth_kb(db, kb_id, current_user, write=True)

    # 读取文件内容
    try:
        content_bytes = await file.read()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"文件读取失败: {str(e)}")

    from app.services.doc_parser import extract_text
    try:
        content = extract_text(file.filename or "unnamed", content_bytes)
    except Exception as e:
        content = content_bytes.decode("utf-8", errors="ignore")

    file_name = file.filename or "unnamed"
    doc_title = title or file_name

    # 保存原始文件到磁盘（用于预览/下载）
    kb_upload_dir = os.path.join(settings.UPLOAD_DIR, "kb", kb_id)
    os.makedirs(kb_upload_dir, exist_ok=True)
    safe_name = f"{_uuid.uuid4().hex[:12]}_{file_name}"
    file_path = os.path.join(kb_upload_dir, safe_name)
    try:
        with open(file_path, "wb") as f:
            f.write(content_bytes)
    except Exception as e:
        logger.warning(f"原始文件保存失败（不影响RAG入库）: {e}")
        file_path = None

    rag_engine = RAGEngine(db_session=db)

    try:
        document = await rag_engine.add_document(
            kb_id=kb_id,
            title=doc_title,
            content=content,
            source_type="file",
            file_path=file_path,
            file_name=file_name,
            file_size=len(content_bytes),
            mime_type=file.content_type,
            meta_data={"ext": (file_name.rsplit(".", 1)[-1].lower() if "." in file_name else "")},
        )
        return document.to_dict()
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"文档处理失败: {str(e)}")


@router.get("/knowledge-bases/{kb_id}/documents/{doc_id}", response_model=DocumentResponse)
async def get_document(
    kb_id: str,
    doc_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    result = await db.execute(
        select(KnowledgeDocument).where(
            and_(
                KnowledgeDocument.id == doc_id,
                KnowledgeDocument.kb_id == kb_id,
            )
        )
    )
    doc = result.scalar_one_or_none()

    if not doc:
        raise HTTPException(status_code=404, detail="文档不存在")

    # 验证权限
    await _auth_kb(db, kb_id, current_user)

    return doc.to_dict()


@router.get("/knowledge-bases/{kb_id}/documents/{doc_id}/download")
async def download_document(
    kb_id: str,
    doc_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """下载知识库中上传的原始文件（仅 source_type=file 且 file_path 存在时可下载）"""
    result = await db.execute(
        select(KnowledgeDocument).where(
            and_(
                KnowledgeDocument.id == doc_id,
                KnowledgeDocument.kb_id == kb_id,
            )
        )
    )
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="文档不存在")
    await _auth_kb(db, kb_id, current_user)

    if not doc.file_path or not os.path.isfile(doc.file_path):
        raise HTTPException(status_code=404, detail="原始文件不可用（可能为文本文档或文件未保存）")

    return FileResponse(
        path=doc.file_path,
        filename=doc.file_name or "download",
        media_type=doc.mime_type or "application/octet-stream",
    )


@router.get("/knowledge-bases/{kb_id}/documents/{doc_id}/preview")
async def preview_document(
    kb_id: str,
    doc_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """在线预览知识库上传的原始文件（内联渲染 PDF/图片等，供前端 iframe 渲染；无原始文件时返回纯文本）"""
    result = await db.execute(
        select(KnowledgeDocument).where(
            and_(
                KnowledgeDocument.id == doc_id,
                KnowledgeDocument.kb_id == kb_id,
            )
        )
    )
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="文档不存在")
    await _auth_kb(db, kb_id, current_user)

    if not doc.file_path or not os.path.isfile(doc.file_path):
        if doc.source_type == "file":
            raise HTTPException(
                status_code=404,
                detail="原始文件不可用（该文档历史仅文本入库，请重新上传原文件以启用预览/下载）",
            )
        text = doc.content or ""
        return Response(content=text.encode("utf-8", "ignore"), media_type="text/plain; charset=utf-8")

    ext = (doc.file_name or "").rsplit(".", 1)[-1].lower() if "." in (doc.file_name or "") else ""
    mt = doc.mime_type or _EXT_MIME.get(ext, "application/octet-stream")
    if ext == "pdf":
        mt = "application/pdf"
    elif ext in ("png", "jpg", "jpeg", "gif", "webp", "bmp", "svg") and not mt.startswith("image/"):
        mt = "image/%s" % ("jpeg" if ext == "jpg" else ext)
    safe_name = doc.file_name or "preview"
    disp = "inline; filename*=UTF-8''" + urllib.parse.quote(safe_name)
    return FileResponse(
        path=doc.file_path,
        media_type=mt,
        filename=safe_name,
        headers={"Content-Disposition": disp},
    )


# 支持的预览/纯文本提取类型
_PREVIEW_TEXT_EXTS = {"txt", "md", "csv", "json", "html", "htm", "xml", "log", "yml", "yaml"}
_PREVIEW_IMAGE_EXTS = {"png", "jpg", "jpeg", "gif", "webp", "bmp", "svg"}
_PREVIEW_PDF_EXTS = {"pdf"}
_PREVIEW_DOCX_EXTS = {"docx"}
_PREVIEW_XLSX_EXTS = {"xlsx", "xls"}


def _render_docx_to_html(path: str) -> str:
    """python-docx → HTML 片段（段落 + 简单表格）。"""
    from docx import Document  # type: ignore
    doc = Document(path)
    out = []
    for p in doc.paragraphs:
        text = (p.text or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        if not text.strip():
            out.append("<p></p>")
            continue
        if p.style and p.style.name and p.style.name.lower().startswith("heading"):
            try:
                level = max(1, min(6, int(p.style.name.split()[-1])))
            except Exception:
                level = 2
            out.append(f"<h{level}>{text}</h{level}>")
        else:
            out.append(f"<p>{text}</p>")
    for tbl in doc.tables:
        out.append('<table style="border-collapse:collapse;width:100%;margin:8px 0;">')
        for row in tbl.rows:
            out.append("<tr>")
            for cell in row.cells:
                txt = (cell.text or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                out.append(
                    f'<td style="border:1px solid #d9d9d9;padding:6px 10px;">{txt}</td>'
                )
            out.append("</tr>")
        out.append("</table>")
    return "".join(out) or "<p style='color:#999;'>（空文档）</p>"


def _render_xlsx_to_html(path: str, sheet_name: Optional[str] = None, max_rows: int = 500) -> dict:
    """openpyxl → HTML 表格片段 + 列出所有 sheet。"""
    import openpyxl  # type: ignore
    wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
    sheets = wb.sheetnames
    target = sheet_name if sheet_name in sheets else (sheets[0] if sheets else None)
    if not target:
        return {"sheets": [], "html": "<p style='color:#999;'>（空工作簿）</p>"}
    ws = wb[target]
    out = [
        f'<h4 style="margin:8px 0 4px;">Sheet: {target}</h4>',
        '<table style="border-collapse:collapse;width:100%;font-size:13px;">',
    ]
    n = 0
    for row in ws.iter_rows(values_only=True):
        if n >= max_rows:
            out.append(
                f'<tr><td colspan="20" style="color:#999;padding:6px;">（仅显示前 {max_rows} 行，共 {ws.max_row} 行）</td></tr>'
            )
            break
        out.append("<tr>")
        for v in row:
            s = "" if v is None else str(v)
            s = s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            out.append(
                f'<td style="border:1px solid #d9d9d9;padding:4px 8px;max-width:240px;white-space:nowrap;overflow:hidden;text-overflow:ellipsis;">{s}</td>'
            )
        out.append("</tr>")
        n += 1
    out.append("</table>")
    return {"sheets": sheets, "html": "".join(out)}


def _extract_pdf_text_pages(path: str, max_pages: int = 30) -> dict:
    """pypdf → 每页纯文本 + 总页数。"""
    from pypdf import PdfReader  # type: ignore
    reader = PdfReader(path)
    total = len(reader.pages)
    pages = []
    for i, p in enumerate(reader.pages[:max_pages]):
        try:
            text = p.extract_text() or ""
        except Exception as e:  # noqa
            text = f"[第 {i+1} 页文本提取失败: {e}]"
        pages.append({"page": i + 1, "text": text})
    return {"total_pages": total, "pages": pages, "truncated": total > max_pages}


@router.get("/knowledge-bases/{kb_id}/documents/{doc_id}/preview.json")
async def preview_document_json(
    kb_id: str,
    doc_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """结构化预览：按文件类型返回 {type, html?, text?, pages?, sheets?, total_pages?}

    type 取值:
    - pdf     → {type, file_url, pages:[{page,text}], total_pages, truncated}
    - docx    → {type, html}
    - xlsx    → {type, sheets, html}
    - image   → {type, file_url, mime}
    - text    → {type, text, mime}
    - other   → {type, file_url, mime}  (前端用 iframe/下载)
    """
    result = await db.execute(
        select(KnowledgeDocument).where(
            and_(KnowledgeDocument.id == doc_id, KnowledgeDocument.kb_id == kb_id)
        )
    )
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="文档不存在")
    await _auth_kb(db, kb_id, current_user)

    if not doc.file_path or not os.path.isfile(doc.file_path):
        # 无原始文件时降级为纯文本
        return {
            "type": "text",
            "mime": "text/plain; charset=utf-8",
            "file_name": doc.file_name or "",
            "text": (doc.content or "")[:200000],
            "warning": "原始文件不可用，仅显示已入库文本",
        }

    ext = (doc.file_name or "").rsplit(".", 1)[-1].lower() if "." in (doc.file_name or "") else ""
    mime = doc.mime_type or _EXT_MIME.get(ext, "application/octet-stream")
    file_url = f"/api/v1/knowledge-bases/{kb_id}/documents/{doc_id}/download"

    if ext in _PREVIEW_PDF_EXTS:
        try:
            data = _extract_pdf_text_pages(doc.file_path)
        except Exception as e:  # noqa
            return {"type": "other", "mime": mime, "file_name": doc.file_name, "file_url": file_url, "error": f"PDF 解析失败: {e}"}
        return {"type": "pdf", "file_name": doc.file_name, "file_url": file_url, **data}

    if ext in _PREVIEW_DOCX_EXTS:
        try:
            html = _render_docx_to_html(doc.file_path)
        except Exception as e:  # noqa
            return {"type": "other", "mime": mime, "file_name": doc.file_name, "file_url": file_url, "error": f"Word 解析失败: {e}"}
        return {"type": "docx", "file_name": doc.file_name, "html": html}

    if ext in _PREVIEW_XLSX_EXTS:
        try:
            data = _render_xlsx_to_html(doc.file_path)
        except Exception as e:  # noqa
            return {"type": "other", "mime": mime, "file_name": doc.file_name, "file_url": file_url, "error": f"Excel 解析失败: {e}"}
        return {"type": "xlsx", "file_name": doc.file_name, **data}

    if ext in _PREVIEW_IMAGE_EXTS:
        return {"type": "image", "file_name": doc.file_name, "file_url": file_url, "mime": mime}

    if ext in _PREVIEW_TEXT_EXTS:
        try:
            with open(doc.file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read(200000)
        except Exception as e:  # noqa
            return {"type": "other", "mime": mime, "file_name": doc.file_name, "file_url": file_url, "error": str(e)}
        return {"type": "text", "file_name": doc.file_name, "text": text, "mime": mime}

    # 其余类型（pptx/rar/zip/未知）：返回下载链接 + 元信息，由前端选择下载或用浏览器原生
    return {
        "type": "other",
        "file_name": doc.file_name,
        "file_url": file_url,
        "mime": mime,
        "ext": ext,
    }


@router.put("/knowledge-bases/{kb_id}/documents/{doc_id}", response_model=DocumentResponse)
async def update_document(
    kb_id: str,
    doc_id: str,
    data: DocumentUpdate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """更新文档标题/正文并重新分块（拆解）+ 异步向量化。

    用于"文档与知识库合并"后富文本文档的在线编辑：编辑即重新入库。
    """
    # 验证知识库写权限
    await _auth_kb(db, kb_id, current_user, write=True)

    rag_engine = RAGEngine(db_session=db)
    try:
        document = await rag_engine.update_document(
            doc_id=doc_id,
            title=data.title,
            content=data.content,
            meta_data=data.meta_data,
        )
        return document.to_dict()
    except RuntimeError as e:
        raise HTTPException(status_code=404, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"文档更新失败: {str(e)}")


@router.delete("/knowledge-bases/documents/{doc_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_document(
    doc_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    # 先获取文档以确认存在并拿到 kb_id
    result = await db.execute(
        select(KnowledgeDocument).where(KnowledgeDocument.id == doc_id)
    )
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="文档不存在或无权访问")

    # 验证知识库写权限（owner / 写分享 / 超管）
    await _auth_kb(db, doc.kb_id, current_user, write=True)

    rag_engine = RAGEngine(db_session=db)
    await rag_engine.delete_document(doc_id)

    return None


# ============== Search Endpoints ==============

@router.post("/knowledge-bases/{kb_id}/search", response_model=SearchResponse)
async def search_knowledge_base(
    kb_id: str,
    data: SearchQuery,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    # 验证知识库权限
    await _auth_kb(db, kb_id, current_user)

    rag_engine = RAGEngine(db_session=db)

    try:
        results = await rag_engine.search(
            kb_id=kb_id,
            query=data.query,
            top_k=data.top_k,
        )

        return SearchResponse(
            query=data.query,
            results=[SearchResult(**r) for r in results],
            total=len(results),
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"搜索失败: {str(e)}")


@router.post("/knowledge-bases/search/multi", response_model=SearchResponse)
async def search_multiple_knowledge_bases(
    kb_ids: List[str],
    data: SearchQuery,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    # 验证所有知识库权限（是否在可访问集合内）
    accessible = set(await get_accessible_kb_ids(db, current_user, scope="all"))
    for kb_id in kb_ids:
        if kb_id not in accessible:
            raise HTTPException(status_code=404, detail=f"知识库 {kb_id} 不存在或无权访问")

    rag_engine = RAGEngine(db_session=db)

    try:
        results = await rag_engine.multi_kb_search(
            kb_ids=kb_ids,
            query=data.query,
            top_k=data.top_k,
        )

        return SearchResponse(
            query=data.query,
            results=[SearchResult(**r) for r in results],
            total=len(results),
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"搜索失败: {str(e)}")


# ============== RAG Context Endpoint ==============

@router.post("/knowledge-bases/rag/context", response_model=RAGResponse)
async def get_rag_context(
    data: RAGQuery,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    # 验证所有知识库权限（是否在可访问集合内）
    accessible = set(await get_accessible_kb_ids(db, current_user, scope="all"))
    for kb_id in data.kb_ids:
        if kb_id not in accessible:
            raise HTTPException(status_code=404, detail=f"知识库 {kb_id} 不存在或无权访问")

    rag_engine = RAGEngine(db_session=db)

    try:
        context = await rag_engine.get_context(
            query=data.query,
            kb_ids=data.kb_ids,
            top_k=data.top_k,
        )

        return RAGResponse(
            query=data.query,
            context=context,
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"获取上下文失败: {str(e)}")


# ============== Document Chunks Endpoint ==============

@router.get("/knowledge-bases/{kb_id}/documents/{doc_id}/chunks")
async def get_document_chunks(
    kb_id: str,
    doc_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    # 验证权限
    await _auth_kb(db, kb_id, current_user)

    rag_engine = RAGEngine(db_session=db)

    try:
        chunks = await rag_engine.get_document_chunks(doc_id)
        return {"chunks": chunks}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"获取片段失败: {str(e)}")


# ============== RAG 问答端点 ==============

class QARequest(BaseModel):
    question: str
    top_k: int = 5
    kb_ids: List[str] = []


@router.post("/knowledge-bases/{kb_id}/qa")
async def knowledge_qa(
    kb_id: str,
    data: QARequest,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    await _auth_kb(db, kb_id, current_user)

    rag_engine = RAGEngine(db_session=db)
    try:
        context = await rag_engine.get_context(data.question, [kb_id], top_k=data.top_k)
        hits = await rag_engine.search(kb_id, data.question, top_k=data.top_k)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"检索失败: {str(e)}")

    answer = await _answer_with_context(data.question, context)
    if not answer:
        answer = "未在知识库中找到相关内容。" + (f"\n\n相关资料片段：\n{context[:600]}" if context else "")
    return {
        "answer": answer,
        "sources": [{"title": h["document_title"], "score": round(h["score"], 3)} for h in hits],
    }


@router.post("/knowledge-bases/qa")
async def knowledge_qa_multi(
    data: QARequest,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    kb_ids = data.kb_ids if hasattr(data, "kb_ids") else []
    # 验证所有知识库权限
    accessible = set(await get_accessible_kb_ids(db, current_user, scope="all"))
    for kb_id in kb_ids:
        if kb_id not in accessible:
            raise HTTPException(status_code=404, detail=f"知识库 {kb_id} 不存在或无权访问")
    rag_engine = RAGEngine(db_session=db)
    try:
        context = await rag_engine.get_context(data.question, kb_ids, top_k=data.top_k)
        hits = await rag_engine.multi_kb_search(kb_ids, data.question, top_k=data.top_k)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"检索失败: {str(e)}")

    answer = await _answer_with_context(data.question, context)
    if not answer:
        answer = "未在知识库中找到相关内容。" + (f"\n\n相关资料片段：\n{context[:600]}" if context else "")
    return {
        "answer": answer,
        "sources": [{"title": h["document_title"], "kb_id": h.get("kb_id"), "score": round(h["score"], 3)} for h in hits],
    }
